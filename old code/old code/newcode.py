# import streamlit as st
# import os
# import time
# import random
# import sqlite3
# from datetime import datetime
# from dotenv import load_dotenv
# from openai import OpenAI, RateLimitError
# import tiktoken
# from fpdf import FPDF
# import base64
# import matplotlib.pyplot as plt
# from PyPDF2 import PdfReader
# from docx import Document as DocxDocument

# # ===== Инициализация БД =====
# def init_db():
#     conn = sqlite3.connect("history.db")
#     conn.execute("""
#         CREATE TABLE IF NOT EXISTS checks (
#             id INTEGER PRIMARY KEY AUTOINCREMENT,
#             check_type TEXT,
#             filename TEXT,
#             report TEXT,
#             similarity TEXT,
#             timestamp TEXT,
#             pdf_path TEXT
#         )
#     """)
#     conn.commit()
#     return conn

# def save_to_db(check_type, filename, report, similarity_list, pdf_path):
#     conn = init_db()
#     similarity_str = ",".join(map(str, similarity_list)) if isinstance(similarity_list, list) else str(similarity_list)
#     conn.execute(
#         "INSERT INTO checks (check_type, filename, report, similarity, timestamp, pdf_path) VALUES (?, ?, ?, ?, ?, ?)",
#         (check_type, filename, report, similarity_str, datetime.now().strftime("%Y-%m-%d %H:%M"), pdf_path)
#     )
#     conn.commit()
#     conn.close()

# def load_history():
#     conn = init_db()
#     rows = conn.execute("SELECT * FROM checks ORDER BY id DESC").fetchall()
#     conn.close()
#     return rows

# # ===== API и конфигурация =====
# load_dotenv()
# OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
# if not OPENAI_API_KEY:
#     st.error("❌ API ключ не найден в .env")
#     st.stop()
# client = OpenAI(api_key=OPENAI_API_KEY)

# # ===== Вспомогательные функции =====
# def extract_text(file) -> str:
#     """Извлечение текста из файлов разных форматов с обработкой ошибок"""
#     try:
#         if file.type == "application/pdf":
#             return "".join(page.extract_text() or "" for page in PdfReader(file).pages)
#         elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#             doc = DocxDocument(file)
#             return "\n".join(p.text for p in doc.paragraphs)
#         elif file.type == "text/plain":
#             return file.read().decode("utf-8")
#     except Exception as e:
#         st.error(f"Ошибка чтения файла: {str(e)}")
#         return ""
#     return ""

# def count_tokens(text: str, model="gpt-4") -> int:
#     """Подсчёт токенов в тексте"""
#     try:
#         encoding = tiktoken.encoding_for_model(model)
#         return len(encoding.encode(text))
#     except:
#         return len(text.split())  # Fallback для неизвестных моделей

# def split_text_into_chunks(text, max_tokens=2000, model="gpt-4"):
#     """Разделение текста на части по количеству токенов"""
#     encoding = tiktoken.encoding_for_model(model)
#     tokens = encoding.encode(text)
#     chunks = []
    
#     for i in range(0, len(tokens), max_tokens):
#         chunk_tokens = tokens[i:i+max_tokens]
#         chunks.append(encoding.decode(chunk_tokens))
    
#     return chunks

# def safe_chat_completion(messages, model="gpt-4", max_tokens=1000, retries=3):
#     """Безопасный запрос к OpenAI с обработкой ограничений"""
#     for attempt in range(retries):
#         try:
#             return client.chat.completions.create(
#                 model=model,
#                 messages=messages,
#                 max_tokens=max_tokens,
#                 temperature=0.7
#             )
#         except RateLimitError:
#             wait_time = 20 * (attempt + 1)
#             st.warning(f"⏳ Лимит запросов. Ждём {wait_time} сек...")
#             time.sleep(wait_time)
#         except Exception as e:
#             st.error(f"Ошибка API: {str(e)}")
#             break
#     return None

# def generate_pdf(report: str, file_name="report.pdf"):
#     """Генерация PDF с поддержкой Unicode"""
#     pdf = FPDF()
#     pdf.add_page()
    
#     # Поддержка Unicode (пробуем разные шрифты)
#     font_added = False
#     font_paths = [
#         "DejaVuSansCondensed.ttf",
#         "arial.ttf",
#         "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
#         "C:/Windows/Fonts/arial.ttf"
#     ]
    
#     for font_path in font_paths:
#         try:
#             pdf.add_font('CustomFont', '', font_path, uni=True)
#             pdf.set_font('CustomFont', '', 12)
#             font_added = True
#             break
#         except:
#             continue
    
#     if not font_added:
#         pdf.set_font("Arial", size=12)  # Последний вариант без Unicode

#     # Обработка текста
#     for line in report.split("\n"):
#         try:
#             if line.startswith("###"):
#                 pdf.set_font(size=14, style='B')
#                 pdf.cell(0, 10, line[3:].strip(), ln=True)
#                 pdf.set_font(size=12, style='')
#             else:
#                 pdf.multi_cell(0, 10, line)
#         except UnicodeEncodeError:
#             # Если не удается кодировать, заменяем проблемные символы
#             cleaned_line = line.encode('utf-8', 'replace').decode('utf-8')
#             pdf.multi_cell(0, 10, cleaned_line)
    
#     path = f"/tmp/{file_name}"
#     pdf.output(path)
#     return path

# def get_binary_file_download_link(file_path, label="📥 Скачать PDF"):
#     """Генерация ссылки для скачивания файла"""
#     with open(file_path, 'rb') as f:
#         data = f.read()
#     b64 = base64.b64encode(data).decode()
#     return f'<a href="data:application/pdf;base64,{b64}" download="{os.path.basename(file_path)}">{label}</a>'

# def draw_similarity_chart(scores):
#     """Визуализация результатов проверки"""
#     fig, ax = plt.subplots(figsize=(10, 4))
#     ax.bar(range(1, len(scores)+1), scores, color='#2e7d32')
#     ax.set_title("График совпадений по частям текста", pad=20)
#     ax.set_xlabel("Часть текста")
#     ax.set_ylabel("Процент совпадения")
#     ax.set_ylim(0, 100)
#     st.pyplot(fig)

# # ===== Стили =====
# st.markdown("""
#     <style>
#         .main-title { 
#             color: #2e7d32; 
#             font-family: 'Segoe UI'; 
#             font-size: 36px; 
#             font-weight: bold;
#             margin-bottom: 20px;
#         }
#         .stTabs [data-baseweb="tab"] { 
#             font-size: 18px; 
#             color: #2e7d32;
#             padding: 10px 20px;
#         }
#         .stButton>button {
#             background-color: #2e7d32;
#             color: white;
#             border: none;
#             border-radius: 6px;
#             padding: 0.5em 1em;
#             font-weight: bold;
#             margin: 5px 0;
#             transition: all 0.3s;
#         }
#         .stButton>button:hover { 
#             background-color: #388e3c; 
#             transform: scale(1.02);
#         }
#         .result-box {
#             background-color: #e8f5e9;
#             padding: 15px;
#             border-radius: 10px;
#             font-size: 16px;
#             line-height: 1.5;
#             margin: 15px 0;
#             border-left: 4px solid #2e7d32;
#         }
#         .file-info {
#             background-color: #f1f8e9;
#             padding: 10px;
#             border-radius: 5px;
#             margin: 10px 0;
#             border: 1px solid #c8e6c9;
#         }
#         .spinner-text {
#             color: #2e7d32;
#             font-weight: bold;
#             margin-top: 10px;
#         }
#         .error-box {
#             background-color: #ffebee;
#             padding: 15px;
#             border-radius: 10px;
#             border-left: 4px solid #f44336;
#             margin: 10px 0;
#         }
#     </style>
# """, unsafe_allow_html=True)

# # ===== Основной интерфейс =====
# st.markdown('<div class="main-title">🧠 AI Content Analyzer Pro</div>', unsafe_allow_html=True)

# tab1, tab2, tab3 = st.tabs(["🔍 Проверка", "📋 Результаты", "🗂 История"])

# with tab1:
#     st.subheader("Выберите тип проверки")
#     check_type = st.radio("", ["🤖 Проверка на ИИ", "📚 Проверка на плагиат"], horizontal=True, label_visibility="collapsed")
    
#     if check_type == "🤖 Проверка на ИИ":
#         st.markdown("""
#             <div style='margin: 15px 0; color: #3a3a3a;'>
#                 Загрузите документ. Система проанализирует текст и определит, был ли он сгенерирован искусственным интеллектом.
#             </div>
#         """, unsafe_allow_html=True)
        
#         ai_file = st.file_uploader("📄 Загрузите документ (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"], key="ai_file")
        
#         if st.button("🔎 Анализировать текст на ИИ", key="ai_check"):
#             if ai_file:
#                 with st.spinner("Анализируем текст..."):
#                     try:
#                         text = extract_text(ai_file)
#                         if not text:
#                             st.error("Не удалось извлечь текст из файла")
#                             st.stop()
                        
#                         chunks = split_text_into_chunks(text)
#                         results = []
                        
#                         progress_bar = st.progress(0)
#                         status_text = st.empty()
                        
#                         for i, chunk in enumerate(chunks):
#                             status_text.markdown(f'<div class="spinner-text">Обработка части {i+1} из {len(chunks)}</div>', unsafe_allow_html=True)
#                             progress_bar.progress((i + 1) / len(chunks))
                            
#                             messages = [
#                                 {"role": "system", "content": "Ты эксперт по определению текста, сгенерированного ИИ. Анализируй текст и дай развернутый ответ с примерами."},
#                                 {"role": "user", "content": f"Был ли этот текст сгенерирован ИИ? Ответь подробно с анализом.\n\n{chunk}"}
#                             ]
                            
#                             result = safe_chat_completion(messages)
#                             if result:
#                                 results.append(f"### Часть {i+1}\n{result.choices[0].message.content}")
                        
#                         full_report = "\n\n".join(results)
#                         pdf_path = generate_pdf(full_report, f"ai_check_{ai_file.name}.pdf")
#                         save_to_db("AI Check", ai_file.name, full_report, [], pdf_path)
                        
#                         st.session_state["report"] = full_report
#                         st.session_state["pdf_path"] = pdf_path
#                         st.success("✅ Анализ завершён!")
#                     except Exception as e:
#                         st.error(f"Ошибка при анализе: {str(e)}")
#             else:
#                 st.warning("Пожалуйста, загрузите документ для анализа.")
    
#     else:  # Проверка на плагиат
#         st.markdown("""
#             <div style='margin: 15px 0; color: #3a3a3a;'>
#                 Загрузите эталонные документы и файл для проверки. Система определит совпадения текста и выдаст детальный отчёт.
#             </div>
#         """, unsafe_allow_html=True)
        
#         ref_files = st.file_uploader("📁 Эталонные документы (.txt, .pdf, .docx)", 
#                                    accept_multiple_files=True, 
#                                    type=["txt", "pdf", "docx"], 
#                                    key="ref_files")
        
#         target_file = st.file_uploader("📄 Документ для анализа", 
#                                      type=["txt", "pdf", "docx"], 
#                                      key="plag_file")
        
#         if st.button("🔍 Проверить на плагиат", key="plag_check"):
#             if ref_files and target_file:
#                 with st.spinner("Анализируем совпадения..."):
#                     try:
#                         ref_texts = [extract_text(f) for f in ref_files]
#                         target_text = extract_text(target_file)
                        
#                         if not target_text:
#                             st.error("Не удалось извлечь текст из анализируемого файла")
#                             st.stop()
                        
#                         target_chunks = split_text_into_chunks(target_text)
                        
#                         report = []
#                         similarity_scores = []
                        
#                         progress_bar = st.progress(0)
#                         status_text = st.empty()
                        
#                         for i, chunk in enumerate(target_chunks):
#                             status_text.markdown(f'<div class="spinner-text">Анализ части {i+1} из {len(target_chunks)}</div>', unsafe_allow_html=True)
#                             progress_bar.progress((i + 1) / len(target_chunks))
                            
#                             chunk_report = []
#                             max_similarity = 0
                            
#                             for j, ref in enumerate(ref_texts):
#                                 if not ref:
#                                     continue
                                
#                                 # Улучшенный алгоритм сравнения
#                                 chunk_words = set(chunk.lower().split())
#                                 ref_words = set(ref.lower().split())
#                                 common_words = chunk_words & ref_words
                                
#                                 similarity = len(common_words) / max(len(chunk_words), 1) * 100
                                
#                                 if similarity > 15:  # Порог для отображения
#                                     chunk_report.append(f"- Совпадение с документом {j+1}: {similarity:.1f}%")
#                                     if similarity > max_similarity:
#                                         max_similarity = similarity
                            
#                             if chunk_report:
#                                 report.append(f"### Часть {i+1}\n" + "\n".join(chunk_report))
#                             else:
#                                 report.append(f"### Часть {i+1}\n- Совпадений не обнаружено")
                            
#                             similarity_scores.append(max_similarity)
                        
#                         full_report = "\n\n".join(report)
#                         pdf_path = generate_pdf(full_report, f"plagiarism_check_{target_file.name}.pdf")
#                         save_to_db("Plagiarism Check", target_file.name, full_report, similarity_scores, pdf_path)
                        
#                         st.session_state["report"] = full_report
#                         st.session_state["similarity"] = similarity_scores
#                         st.session_state["pdf_path"] = pdf_path
#                         st.success("✅ Проверка завершена!")
#                     except Exception as e:
#                         st.error(f"Ошибка при проверке на плагиат: {str(e)}")
#             else:
#                 st.warning("Пожалуйста, загрузите все необходимые файлы.")

# with tab2:
#     if "report" in st.session_state:
#         st.markdown('<div class="result-box">{}</div>'.format(
#             st.session_state["report"].replace("\n", "<br>")), unsafe_allow_html=True)
        
#         if "similarity" in st.session_state and st.session_state["similarity"]:
#             st.markdown("---")
#             st.subheader("Визуализация результатов")
#             draw_similarity_chart(st.session_state["similarity"])
        
#         st.markdown("---")
#         st.markdown(get_binary_file_download_link(
#             st.session_state["pdf_path"], 
#             "📥 Скачать полный отчёт в PDF"
#         ), unsafe_allow_html=True)
#     else:
#         st.info("ℹ️ Запустите проверку во вкладке '🔍 Проверка'.")

# with tab3:
#     st.subheader("История проверок")
#     history = load_history()
    
#     if history:
#         for row in history:
#             with st.expander(f"{row[2]} — {row[1]} ({row[5]})"):
#                 st.markdown(f"**Тип проверки:** {row[1]}")
#                 st.markdown(f"**Файл:** {row[2]}")
#                 st.markdown(f"**Дата:** {row[5]}")
                
#                 if st.button(f"Показать отчёт #{row[0]}", key=f"show_{row[0]}"):
#                     st.markdown('<div class="result-box">{}</div>'.format(
#                         row[3].replace("\n", "<br>")), unsafe_allow_html=True)
                    
#                     if row[4]:
#                         try:
#                             scores = list(map(float, row[4].split(",")))
#                             draw_similarity_chart(scores)
#                         except:
#                             pass
                    
#                     st.markdown(get_binary_file_download_link(
#                         row[6], 
#                         f"📥 Скачать отчёт #{row[0]}"
#                     ), unsafe_allow_html=True)
                
#                 st.markdown("---")
#     else:
#         st.info("📭 История проверок пуста. Загрузите файл для начала анализа.")


import streamlit as st
import os
import time
import random
import sqlite3
from datetime import datetime
from dotenv import load_dotenv
from openai import OpenAI, RateLimitError
import tiktoken
from fpdf import FPDF
import base64
import matplotlib.pyplot as plt
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# ===== Инициализация БД =====
def init_db():
    conn = sqlite3.connect("history.db")
    conn.execute("""
        CREATE TABLE IF NOT EXISTS checks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            check_type TEXT,
            filename TEXT,
            report TEXT,
            similarity TEXT,
            timestamp TEXT,
            pdf_path TEXT
        )
    """)
    conn.commit()
    return conn

def save_to_db(check_type, filename, report, similarity_list, pdf_path):
    conn = init_db()
    similarity_str = ",".join(map(str, similarity_list)) if isinstance(similarity_list, list) else str(similarity_list)
    conn.execute(
        "INSERT INTO checks (check_type, filename, report, similarity, timestamp, pdf_path) VALUES (?, ?, ?, ?, ?, ?)",
        (check_type, filename, report, similarity_str, datetime.now().strftime("%Y-%m-%d %H:%M"), pdf_path)
    )
    conn.commit()
    conn.close()

def load_history():
    conn = init_db()
    rows = conn.execute("SELECT * FROM checks ORDER BY id DESC").fetchall()
    conn.close()
    return rows

# ===== API и конфигурация =====
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    st.error("❌ API ключ не найден в .env")
    st.stop()
client = OpenAI(api_key=OPENAI_API_KEY)

# ===== Вспомогательные функции =====
def extract_text(file) -> str:
    """Извлечение текста из файлов разных форматов с обработкой ошибок"""
    try:
        if file.type == "application/pdf":
            return "".join(page.extract_text() or "" for page in PdfReader(file).pages)
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(file)
            return "\n".join(p.text for p in doc.paragraphs)
        elif file.type == "text/plain":
            return file.read().decode("utf-8")
    except Exception as e:
        st.error(f"Ошибка чтения файла: {str(e)}")
        return ""
    return ""

def count_tokens(text: str, model="gpt-4") -> int:
    """Подсчёт токенов в тексте"""
    try:
        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))
    except:
        return len(text.split())  # Fallback для неизвестных моделей

def split_text_into_chunks(text, max_tokens=2000, model="gpt-4"):
    """Разделение текста на части по количеству токенов"""
    encoding = tiktoken.encoding_for_model(model)
    tokens = encoding.encode(text)
    chunks = []
    
    for i in range(0, len(tokens), max_tokens):
        chunk_tokens = tokens[i:i+max_tokens]
        chunks.append(encoding.decode(chunk_tokens))
    
    return chunks

def safe_chat_completion(messages, model="gpt-4", max_tokens=1000, retries=3):
    """Безопасный запрос к OpenAI с обработкой ограничений"""
    for attempt in range(retries):
        try:
            return client.chat.completions.create(
                model=model,
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.7
            )
        except RateLimitError:
            wait_time = 20 * (attempt + 1)
            st.warning(f"⏳ Лимит запросов. Ждём {wait_time} сек...")
            time.sleep(wait_time)
        except Exception as e:
            st.error(f"Ошибка API: {str(e)}")
            break
    return None

def generate_pdf(report: str, file_name="report.pdf"):
    """Генерация PDF с поддержкой Unicode"""
    pdf = FPDF()
    pdf.add_page()
    
    # Поддержка Unicode (пробуем разные шрифты)
    font_added = False
    current_font = "Arial"  # По умолчанию
    
    font_paths = [
        "DejaVuSansCondensed.ttf",
        "arial.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arial.ttf"
    ]
    
    for font_path in font_paths:
        try:
            pdf.add_font('CustomFont', '', font_path, uni=True)
            current_font = 'CustomFont'
            font_added = True
            break
        except:
            continue
    
    # Устанавливаем шрифт
    pdf.set_font(current_font, size=12)

    # Обработка текста
    for line in report.split("\n"):
        try:
            if line.startswith("###"):
                pdf.set_font(current_font, size=14, style='B')
                pdf.cell(0, 10, line[3:].strip(), ln=True)
                pdf.set_font(current_font, size=12, style='')
            else:
                pdf.multi_cell(0, 10, line)
        except UnicodeEncodeError:
            # Если не удается кодировать, заменяем проблемные символы
            cleaned_line = line.encode('utf-8', 'replace').decode('utf-8')
            pdf.multi_cell(0, 10, cleaned_line)
    
    path = f"/tmp/{file_name}"
    pdf.output(path)
    return path

def get_binary_file_download_link(file_path, label="📥 Скачать PDF"):
    """Генерация ссылки для скачивания файла"""
    with open(file_path, 'rb') as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    return f'<a href="data:application/pdf;base64,{b64}" download="{os.path.basename(file_path)}">{label}</a>'

def draw_similarity_chart(scores):
    """Визуализация результатов проверки"""
    fig, ax = plt.subplots(figsize=(10, 4))
    ax.bar(range(1, len(scores)+1), scores, color='#2e7d32')
    ax.set_title("График совпадений по частям текста", pad=20)
    ax.set_xlabel("Часть текста")
    ax.set_ylabel("Процент совпадения")
    ax.set_ylim(0, 100)
    st.pyplot(fig)

# ===== Стили =====
st.markdown("""
    <style>
        .main-title { 
            color: #2e7d32; 
            font-family: 'Segoe UI'; 
            font-size: 36px; 
            font-weight: bold;
            margin-bottom: 20px;
        }
        .stTabs [data-baseweb="tab"] { 
            font-size: 18px; 
            color: #2e7d32;
            padding: 10px 20px;
        }
        .stButton>button {
            background-color: #2e7d32;
            color: white;
            border: none;
            border-radius: 6px;
            padding: 0.5em 1em;
            font-weight: bold;
            margin: 5px 0;
            transition: all 0.3s;
        }
        .stButton>button:hover { 
            background-color: #388e3c; 
            transform: scale(1.02);
        }
        .result-box {
            background-color: #e8f5e9;
            padding: 15px;
            border-radius: 10px;
            font-size: 16px;
            line-height: 1.5;
            margin: 15px 0;
            border-left: 4px solid #2e7d32;
        }
        .file-info {
            background-color: #f1f8e9;
            padding: 10px;
            border-radius: 5px;
            margin: 10px 0;
            border: 1px solid #c8e6c9;
        }
        .spinner-text {
            color: #2e7d32;
            font-weight: bold;
            margin-top: 10px;
        }
        .error-box {
            background-color: #ffebee;
            padding: 15px;
            border-radius: 10px;
            border-left: 4px solid #f44336;
            margin: 10px 0;
        }
    </style>
""", unsafe_allow_html=True)

# ===== Основной интерфейс =====
st.markdown('<div class="main-title">🧠 AI Content Analyzer Pro</div>', unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs(["🔍 Проверка", "📋 Результаты", "🗂 История"])

with tab1:
    st.subheader("Выберите тип проверки")
    check_type = st.radio("", ["🤖 Проверка на ИИ", "📚 Проверка на плагиат"], horizontal=True, label_visibility="collapsed")
    
    if check_type == "🤖 Проверка на ИИ":
        st.markdown("""
            <div style='margin: 15px 0; color: #3a3a3a;'>
                Загрузите документ. Система проанализирует текст и определит, был ли он сгенерирован искусственным интеллектом.
            </div>
        """, unsafe_allow_html=True)
        
        ai_file = st.file_uploader("📄 Загрузите документ (.txt, .pdf, .docx)", type=["txt", "pdf", "docx"], key="ai_file")
        
        if st.button("🔎 Анализировать текст на ИИ", key="ai_check"):
            if ai_file:
                with st.spinner("Анализируем текст..."):
                    try:
                        text = extract_text(ai_file)
                        if not text:
                            st.error("Не удалось извлечь текст из файла")
                            st.stop()
                        
                        chunks = split_text_into_chunks(text)
                        results = []
                        
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        for i, chunk in enumerate(chunks):
                            status_text.markdown(f'<div class="spinner-text">Обработка части {i+1} из {len(chunks)}</div>', unsafe_allow_html=True)
                            progress_bar.progress((i + 1) / len(chunks))
                            
                            messages = [
                                {"role": "system", "content": "Ты эксперт по определению текста, сгенерированного ИИ. Анализируй текст и дай развернутый ответ с примерами."},
                                {"role": "user", "content": f"Был ли этот текст сгенерирован ИИ? Ответь подробно с анализом.\n\n{chunk}"}
                            ]
                            
                            result = safe_chat_completion(messages)
                            if result:
                                results.append(f"### Часть {i+1}\n{result.choices[0].message.content}")
                        
                        full_report = "\n\n".join(results)
                        pdf_path = generate_pdf(full_report, f"ai_check_{ai_file.name}.pdf")
                        save_to_db("AI Check", ai_file.name, full_report, [], pdf_path)
                        
                        st.session_state["report"] = full_report
                        st.session_state["pdf_path"] = pdf_path
                        st.success("✅ Анализ завершён!")
                    except Exception as e:
                        st.error(f"Ошибка при анализе: {str(e)}")
            else:
                st.warning("Пожалуйста, загрузите документ для анализа.")
    
    else:  # Проверка на плагиат
        st.markdown("""
            <div style='margin: 15px 0; color: #3a3a3a;'>
                Загрузите эталонные документы и файл для проверки. Система определит совпадения текста и выдаст детальный отчёт.
            </div>
        """, unsafe_allow_html=True)
        
        ref_files = st.file_uploader("📁 Эталонные документы (.txt, .pdf, .docx)", 
                                   accept_multiple_files=True, 
                                   type=["txt", "pdf", "docx"], 
                                   key="ref_files")
        
        target_file = st.file_uploader("📄 Документ для анализа", 
                                     type=["txt", "pdf", "docx"], 
                                     key="plag_file")
        
        if st.button("🔍 Проверить на плагиат", key="plag_check"):
            if ref_files and target_file:
                with st.spinner("Анализируем совпадения..."):
                    try:
                        ref_texts = [extract_text(f) for f in ref_files]
                        target_text = extract_text(target_file)
                        
                        if not target_text:
                            st.error("Не удалось извлечь текст из анализируемого файла")
                            st.stop()
                        
                        target_chunks = split_text_into_chunks(target_text)
                        
                        report = []
                        similarity_scores = []
                        
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        for i, chunk in enumerate(target_chunks):
                            status_text.markdown(f'<div class="spinner-text">Анализ части {i+1} из {len(target_chunks)}</div>', unsafe_allow_html=True)
                            progress_bar.progress((i + 1) / len(target_chunks))
                            
                            chunk_report = []
                            max_similarity = 0
                            
                            for j, ref in enumerate(ref_texts):
                                if not ref:
                                    continue
                                
                                # Улучшенный алгоритм сравнения
                                chunk_words = set(chunk.lower().split())
                                ref_words = set(ref.lower().split())
                                common_words = chunk_words & ref_words
                                
                                similarity = len(common_words) / max(len(chunk_words), 1) * 100
                                
                                if similarity > 15:  # Порог для отображения
                                    chunk_report.append(f"- Совпадение с документом {j+1}: {similarity:.1f}%")
                                    if similarity > max_similarity:
                                        max_similarity = similarity
                            
                            if chunk_report:
                                report.append(f"### Часть {i+1}\n" + "\n".join(chunk_report))
                            else:
                                report.append(f"### Часть {i+1}\n- Совпадений не обнаружено")
                            
                            similarity_scores.append(max_similarity)
                        
                        full_report = "\n\n".join(report)
                        pdf_path = generate_pdf(full_report, f"plagiarism_check_{target_file.name}.pdf")
                        save_to_db("Plagiarism Check", target_file.name, full_report, similarity_scores, pdf_path)
                        
                        st.session_state["report"] = full_report
                        st.session_state["similarity"] = similarity_scores
                        st.session_state["pdf_path"] = pdf_path
                        st.success("✅ Проверка завершена!")
                    except Exception as e:
                        st.error(f"Ошибка при проверке на плагиат: {str(e)}")
            else:
                st.warning("Пожалуйста, загрузите все необходимые файлы.")

with tab2:
    if "report" in st.session_state:
        st.markdown('<div class="result-box">{}</div>'.format(
            st.session_state["report"].replace("\n", "<br>")), unsafe_allow_html=True)
        
        if "similarity" in st.session_state and st.session_state["similarity"]:
            st.markdown("---")
            st.subheader("Визуализация результатов")
            draw_similarity_chart(st.session_state["similarity"])
        
        st.markdown("---")
        st.markdown(get_binary_file_download_link(
            st.session_state["pdf_path"], 
            "📥 Скачать полный отчёт в PDF"
        ), unsafe_allow_html=True)
    else:
        st.info("ℹ️ Запустите проверку во вкладке '🔍 Проверка'.")

with tab3:
    st.subheader("История проверок")
    history = load_history()
    
    if history:
        for row in history:
            with st.expander(f"{row[2]} — {row[1]} ({row[5]})"):
                st.markdown(f"**Тип проверки:** {row[1]}")
                st.markdown(f"**Файл:** {row[2]}")
                st.markdown(f"**Дата:** {row[5]}")
                
                if st.button(f"Показать отчёт #{row[0]}", key=f"show_{row[0]}"):
                    st.markdown('<div class="result-box">{}</div>'.format(
                        row[3].replace("\n", "<br>")), unsafe_allow_html=True)
                    
                    if row[4]:
                        try:
                            scores = list(map(float, row[4].split(",")))
                            draw_similarity_chart(scores)
                        except:
                            pass
                    
                    st.markdown(get_binary_file_download_link(
                        row[6], 
                        f"📥 Скачать отчёт #{row[0]}"
                    ), unsafe_allow_html=True)
                
                st.markdown("---")
    else:
        st.info("📭 История проверок пуста. Загрузите файл для начала анализа.")
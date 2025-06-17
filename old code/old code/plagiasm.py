import streamlit as st
import os
from dotenv import load_dotenv

# OpenAI & LangChain imports
import openai
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.llms import OpenAI

# File parsers
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

# --- Load environment variables ---
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY not found in .env")
openai.api_key = OPENAI_API_KEY

# --- Extract text from supported files ---
def extract_text(file) -> str:
    if file.type == "application/pdf":
        reader = PdfReader(file)
        return "".join(page.extract_text() or "" for page in reader.pages)

    elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = DocxDocument(file)
        return "\n".join(p.text for p in doc.paragraphs)

    elif file.type == "text/plain":
        return file.read().decode("utf-8")

    else:
        return "❌ Формат файла не поддерживается."

# --- Build and persist Chroma vector store ---
@st.cache_resource(show_spinner=False)
def build_vector_store(texts: list[str], persist_dir: str = "./chroma_store") -> Chroma:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    docs = []
    for i, txt in enumerate(texts):
        chunks = splitter.split_text(txt)
        docs.extend([Document(page_content=chunk, metadata={"source": f"doc_{i}"}) for chunk in chunks])

    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    vectordb = Chroma.from_documents(docs, embeddings, persist_directory=persist_dir)
    vectordb.persist()
    return vectordb

# --- Load vector store and set up retrieval QA ---
@st.cache_resource(show_spinner=False)
def get_retriever(k: int = 5) -> RetrievalQA:
    embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
    vectordb = Chroma(persist_directory="./chroma_store", embedding_function=embeddings)
    retriever = vectordb.as_retriever(search_kwargs={"k": k})
    llm = OpenAI(temperature=0, max_tokens=500, openai_api_key=OPENAI_API_KEY)
    return RetrievalQA.from_chain_type(llm, chain_type="stuff", retriever=retriever)

# --- Streamlit UI ---
st.set_page_config(page_title="RAG Plagiarism Checker", layout="centered")
st.title("🔍 Проверка на плагиат с помощью GPT + LangChain")
st.write("Загрузите эталонные документы, чтобы построить индекс, затем проверьте новые тексты на плагиат.")

# Step 1: Upload reference documents
ref_files = st.file_uploader(
    "📚 Загрузите эталонные документы (.txt, .pdf, .docx)",
    accept_multiple_files=True,
    type=["txt", "pdf", "docx"]
)
if ref_files:
    if st.button("🔧 Построить индекс"):
        with st.spinner("Идёт обработка и построение индекса..."):
            texts = [extract_text(f) for f in ref_files]
            build_vector_store(texts)
        st.success("Индекс успешно построен!")

# Step 2: Upload file to check
uploaded_file = st.file_uploader(
    "📄 Загрузите файл для проверки на плагиат",
    type=["txt", "pdf", "docx"],
    key="check"
)
if uploaded_file:
    full_text = extract_text(uploaded_file)
    text_to_check = full_text[:3000]  # ограничиваем размер
    st.subheader("📌 Фрагмент для анализа")
    st.text_area("", text_to_check[:1000], height=200)

    if st.button("🚨 Проверить на плагиат"):
        with st.spinner("Выполняется анализ..."):
            qa = get_retriever()
            query = (
                "Оцени, является ли следующий текст плагиатом относительно эталонных документов. "
                "Если да, укажи похожие фрагменты и вероятность плагиата в процентах (0-100%).\n\n"
                f"Текст:\n{text_to_check}"
            )
            result = qa.run(query)
        st.subheader("📊 Результат анализа")
        st.write(result)

# Footer
st.markdown("---")
st.markdown("🔧 Сделано с ❤️ на основе Streamlit, LangChain и OpenAI GPT")
